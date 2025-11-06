from flask import Blueprint, request, jsonify, send_file, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from datetime import datetime
from app.models import db
from app.models.document import Document, DocumentVersion, DocumentCategory as Category
from app.models.access_log import AccessLog
from app.utils.auth import verify_permission, get_current_user, check_document_permission
from app.utils.file_handler import get_file_type, save_uploaded_file, delete_file, get_file_path, check_file_size, get_file_size
from app.utils.limiter import check_upload_limit
from app.services.log_service import LogService
import docx

# 创建蓝图
documents_bp = Blueprint('documents', __name__)


@documents_bp.route('/', methods=['GET'])
@jwt_required()
@verify_permission('view')
def get_documents():
    """获取文档列表"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 分页参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        
        # 搜索和筛选参数
        keyword = request.args.get('keyword')
        category_id = request.args.get('category_id', type=int)
        file_type = request.args.get('file_type')
        is_my_documents = request.args.get('is_my_documents', type=bool)
        
        # 构建查询
        query = Document.query
        
        # 如果不是管理员，只能看到自己的文档和公开文档
        if user.role.name != 'admin':
            query = query.filter(
                (Document.creator_id == user.id) | (Document.is_private == False)
            )
        
        # 如果指定了只看自己的文档
        if is_my_documents:
            query = query.filter(Document.creator_id == user.id)
        
        # 关键词搜索
        if keyword:
            query = query.filter(
                (Document.title.like(f'%{keyword}%')) | (Document.description.like(f'%{keyword}%'))
            )
        
        # 分类筛选
        if category_id:
            query = query.filter(Document.category_id == category_id)
        
        # 文件类型筛选
        if file_type:
            query = query.filter(Document.file_type == file_type)
        
        # 执行查询
        pagination = query.order_by(Document.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        # 构建响应
        documents = []
        for doc in pagination.items:
            documents.append({
                'id': doc.id,
                'title': doc.title,
                'description': doc.description,
                'file_path': doc.file_path,
                'file_name': doc.file_name,
                'file_type': doc.file_type,
                'category_id': doc.category_id,
                'category_name': doc.category.name if doc.category else None,
                'user_id': doc.creator_id,
                'username': doc.creator.username if doc.creator else '',
                'is_private': doc.is_private,
                'file_size': doc.file_size,
                'version': 1,  # 默认版本号
                'created_at': doc.created_at.isoformat(),
                'updated_at': doc.updated_at.isoformat()
            })
        
        return jsonify({
            'documents': documents,
            'total': pagination.total,
            'page': page,
            'per_page': per_page
        })
    
    except Exception as e:
        return jsonify({'message': f'获取文档列表失败: {str(e)}'}), 500


@documents_bp.route('/', methods=['POST'])
@jwt_required()
@verify_permission('upload')
def upload_document():
    """上传文档"""
    try:
        # 添加详细日志
        print(f"[DEBUG] 收到上传请求，表单数据: {dict(request.form)}")
        print(f"[DEBUG] 请求头: {dict(request.headers)}")
        
        # 获取当前用户
        user = get_current_user()
        print(f"[DEBUG] 当前用户: {user.id}, {user.username}")
        
        # 检查上传限制
        print(f"[DEBUG] 检查上传限制，用户ID: {user.id}")
        if not check_upload_limit(user.id):
            return jsonify({'message': '今日上传文件数量已达上限'}), 403
        
        # 获取上传的文件
        file = request.files.get('file')
        print(f"[DEBUG] 获取文件: {file.filename if file else None}")
        if not file or file.filename == '':
            return jsonify({'message': '请选择要上传的文件'}), 400
        
        # 检查文件类型
        file_type = get_file_type(file.filename)
        print(f"[DEBUG] 文件类型: {file_type}, 文件名: {file.filename}")
        if not file_type:
            return jsonify({'message': '不支持的文件格式'}), 400
        
        # 详细日志记录文件大小信息
        print(f"[DEBUG] 文件对象属性检查 - content_length: {file.content_length}")
        if hasattr(file, 'mimetype'):
            print(f"[DEBUG] 文件MIME类型: {file.mimetype}")
        
        # 尝试获取文件大小的多种方式
        content_length = file.content_length
        print(f"[DEBUG] 从content_length获取文件大小: {content_length} 字节")
        
        # 检查文件大小是否有效
        if content_length is None:
            print(f"[WARNING] content_length为None，无法直接获取文件大小")
        elif content_length <= 0:
            print(f"[WARNING] content_length为0或负数: {content_length}")
        
        # 检查文件大小
        print(f"[DEBUG] 检查文件大小: {content_length or 0} 字节")
        if not check_file_size(file):
            return jsonify({'message': '文件大小超过限制'}), 400
        
        # 保存文件
        print(f"[DEBUG] 开始保存文件，类型: {file_type}")
        file_path, unique_filename = save_uploaded_file(file, file_type)
        print(f"[DEBUG] 文件保存成功，路径: {file_path}, 唯一文件名: {unique_filename}")
        
        # 使用新的get_file_size函数获取文件大小
        print(f"[DEBUG] 调用get_file_size函数获取文件大小，路径: {file_path}")
        file_size = get_file_size(file_path)
        print(f"[DEBUG] 使用get_file_size函数获取文件大小: {file_size} 字节")
        
        # 对比content_length和实际大小
        if content_length is not None and content_length != file_size:
            print(f"[WARNING] 文件大小不匹配 - content_length: {content_length}, 实际大小: {file_size}")
        
        print(f"[DEBUG] 最终保存到数据库的文件大小: {file_size} 字节")
        
        # 获取表单数据
        title = request.form.get('title', file.filename)
        description = request.form.get('description', '')
        category_id = request.form.get('category_id', type=int)
        is_private = request.form.get('is_private', 'false').lower() == 'true'
        print(f"[DEBUG] 表单数据: title={title}, category_id={category_id}, is_private={is_private}")
        
        # 创建文档记录
        document = Document(
            title=title,
            description=description,
            file_path=file_path,
            file_name=unique_filename,
            file_type=file_type,
            category_id=category_id,
            creator_id=user.id,
            is_private=is_private,
            file_size=file_size,
            document_type=file_type  # 使用文件类型作为文档类型
        )
        
        db.session.add(document)
        db.session.commit()
        
        # 记录访问日志
        log = AccessLog(
            user_id=user.id,
            document_id=document.id,
            action_type='upload',
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )
        db.session.add(log)
        db.session.commit()
        
        return jsonify({'message': '文档上传成功', 'document_id': document.id}), 201
    
    except Exception as e:
        db.session.rollback()
        print(f"[ERROR] 上传失败: {str(e)}")
        import traceback
        print(f"[ERROR] 错误堆栈: {traceback.format_exc()}")
        return jsonify({'message': f'文档上传失败: {str(e)}'}), 500


@documents_bp.route('/<int:document_id>', methods=['GET'])
@jwt_required()
@verify_permission('view')
def get_document(document_id):
    """获取文档详情"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if not check_document_permission(user, document):
            return jsonify({'message': '无权限访问此文档'}), 403
        
        # 记录文档访问日志
        LogService.log_document_access(user, document, request)
        
        # 记录访问日志到数据库
        log = AccessLog(
            user_id=user.id,
            document_id=document.id,
            action_type='view',
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )
        db.session.add(log)
        db.session.commit()
        
        # 返回文档信息
        return jsonify({
            'document': {
                'id': document.id,
                'title': document.title,
                'description': document.description,
                'file_path': document.file_path,
                'file_name': document.file_name,
                'original_filename': document.file_name,  # 使用file_name代替不存在的original_filename
                'file_type': document.file_type,
                'category_id': document.category_id,
                'category_name': document.category.name if document.category else None,
                'creator_id': document.creator_id,
                'username': document.creator.username if document.creator else '',
                'is_private': document.is_private,
                'file_size': document.file_size,
                'version': 1,  # 默认版本号,
                'content': document.content,  # 仅流式文件有内容
                'created_at': document.created_at.isoformat(),
                'updated_at': document.updated_at.isoformat()
            }
        })
    
    except Exception as e:
        return jsonify({'message': f'获取文档详情失败: {str(e)}'}), 500


@documents_bp.route('/<int:document_id>/preview', methods=['GET'])
@jwt_required()
@verify_permission('view')
def preview_document(document_id):
    """预览文档内容"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if not check_document_permission(user, document):
            return jsonify({'message': '无权限访问此文档'}), 403
        
        # 获取文件路径
        file_path = get_file_path(document.file_path)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return jsonify({'message': '文件不存在'}), 404
        
        # 记录访问日志
        log = AccessLog(
            user_id=user.id,
            document_id=document.id,
            action_type='preview',
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )
        db.session.add(log)
        db.session.commit()
        
        # 增加查看次数
        document.views_count += 1
        db.session.commit()
        
        # 根据文件扩展名决定返回方式
        _, ext = os.path.splitext(document.file_name.lower())
        
        # 对于Word文档(.docx)，使用python-docx库解析
        if ext == '.docx':
            try:
                doc = docx.Document(file_path)
                content = []
                # 提取文档中所有段落文本
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                
                # 提取表格内容
                for table in doc.tables:
                    table_content = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text)
                        if row_text:
                            table_content.append(' | '.join(row_text))
                    if table_content:
                        content.append('\n表格内容:\n' + '\n'.join(table_content))
                
                # 合并所有内容，用换行符分隔
                full_content = '\n\n'.join(content)
                
                return jsonify({
                    'content': full_content,
                    'file_extension': ext,
                    'file_name': document.file_name
                })
            except Exception as e:
                print(f"解析Word文档失败: {e}")
                # 如果解析失败，继续处理
        
        # 对于文本类型的文件，可以直接返回内容
        if ext in ['.txt', '.md', '.json', '.log', '.csv', '.xml', '.html', '.htm']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return jsonify({
                    'content': content,
                    'file_extension': ext,
                    'file_name': document.file_name
                })
            except UnicodeDecodeError:
                # 如果utf-8解码失败，尝试其他编码
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        content = f.read()
                    return jsonify({
                        'content': content,
                        'file_extension': ext,
                        'file_name': document.file_name
                    })
                except:
                    # 如果都失败，返回文件URL让前端处理
                    pass
        
        # 对于图片类型的文件，返回预览URL
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            # 对于图片，我们直接返回图片的下载URL作为预览URL
            return jsonify({
                'preview_url': f'/api/documents/{document_id}/download',
                'file_extension': ext,
                'file_name': document.file_name,
                'is_image': True
            })
        
        # 对于其他文件类型，返回文件URL让前端通过下载方式处理
        # 生成一个临时的预览URL或直接使用下载端点
        return jsonify({
            'preview_url': f'/api/documents/{document_id}/download',
            'file_extension': ext,
            'file_name': document.file_name,
            'needs_download': True
        })
    
    except Exception as e:
        return jsonify({'message': f'预览文档失败: {str(e)}'}), 500

@documents_bp.route('/<int:document_id>/download', methods=['GET'])
@jwt_required()
@verify_permission('view')
def download_document(document_id):
    """下载文档"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if not check_document_permission(user, document):
            return jsonify({'message': '无权限访问此文档'}), 403
        
        # 获取文件路径
        file_path = get_file_path(document.file_path)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return jsonify({'message': '文件不存在'}), 404
        
        # 记录访问日志
        log = AccessLog(
            user_id=user.id,
            document_id=document.id,
            action_type='download',
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )
        db.session.add(log)
        db.session.commit()
        
        # 返回文件
        return send_file(file_path, as_attachment=True, download_name=document.original_filename)
    
    except Exception as e:
        return jsonify({'message': f'下载文档失败: {str(e)}'}), 500


@documents_bp.route('/<int:document_id>', methods=['PUT'])
@jwt_required()
@verify_permission('edit')
def update_document(document_id):
    """更新文档信息"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if document.creator_id != user.id and user.role.name != 'admin':
            return jsonify({'message': '无权限修改此文档'}), 403
        
        # 检查是否有文件上传
        file = request.files.get('file')
        if file and file.filename != '':
            # 检查文件类型
            file_type = get_file_type(file.filename)
            if not file_type:
                return jsonify({'message': '不支持的文件格式'}), 400
            
            # 检查文件大小
            if not check_file_size(file):
                return jsonify({'message': '文件大小超过限制'}), 400
            
            # 更新文件
            current_app.logger.debug(f"[DEBUG] 开始更新文件，文档ID: {document_id}")
            existing_path = document.file_path
            file_path, unique_filename, file_size = update_uploaded_file(file, file_type, existing_path)
            
            # 更新文档信息
            document.file_path = file_path
            document.file_name = unique_filename
            document.file_type = file_type
            document.file_size = file_size
            current_app.logger.debug(f"[DEBUG] 文件更新成功，新路径: {file_path}, 大小: {file_size} 字节")
        
        # 获取请求数据 (支持表单和JSON两种格式)
        data = {}
        if request.is_json:
            data = request.get_json()
        else:
            data = dict(request.form)
        
        # 更新文档信息
        if 'title' in data:
            document.title = data['title']
        if 'description' in data:
            document.description = data['description']
        if 'category_id' in data:
            document.category_id = int(data['category_id'])
        if 'is_private' in data:
            document.is_private = data['is_private'].lower() == 'true' if isinstance(data['is_private'], str) else bool(data['is_private'])
        
        # 如果是流式文件，可以更新内容
        if document.file_type == 'stream' and 'content' in data:
            # 创建版本记录
            version = DocumentVersion(
                document_id=document.id,
                version_num=1,  # 使用正确的字段名
                content=document.content,
                created_by=user.id,  # 添加缺少的字段
                created_at=document.updated_at
            )
            db.session.add(version)
            
            # 更新内容和版本
            document.content = data['content']
            document.version += 1
        
        document.updated_at = datetime.utcnow()
        db.session.commit()
        
        # 记录访问日志
        log = AccessLog(
            user_id=user.id,
            document_id=document.id,
            action_type='edit',
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )
        db.session.add(log)
        db.session.commit()
        
        return jsonify({'message': '文档更新成功'})
    
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"[ERROR] 更新文档失败: {str(e)}")
        import traceback
        current_app.logger.error(f"[ERROR] 错误堆栈: {traceback.format_exc()}")
        return jsonify({'message': f'更新文档失败: {str(e)}'}), 500


@documents_bp.route('/<int:document_id>', methods=['DELETE'])
@jwt_required()
def delete_document(document_id):
    """删除文档"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if document.creator_id != user.id and user.role.name != 'admin':
            return jsonify({'message': '无权限删除此文档'}), 403
        
        # 首先删除所有相关记录，避免外键约束错误
        # 删除标注记录
        from app.models.annotation import Annotation
        Annotation.query.filter_by(document_id=document_id).delete()
        
        # 删除收藏记录
        from app.models.user_favorite import UserFavorite
        UserFavorite.query.filter_by(document_id=document_id).delete()
        
        # 删除访问日志记录（已在文件顶部导入）
        AccessLog.query.filter_by(document_id=document_id).delete()
        
        # 删除版本记录
        DocumentVersion.query.filter_by(document_id=document_id).delete()
        
        # 删除文件
        delete_file(document.file_path)
        
        # 删除文档
        db.session.delete(document)
        db.session.commit()
        
        return jsonify({'message': '文档删除成功'})
    
    except Exception as e:
        db.session.rollback()
        print(f"删除文档时发生错误: {str(e)}")
        return jsonify({'message': f'删除文档失败: {str(e)}'}), 500


@documents_bp.route('/<int:document_id>/versions', methods=['GET'])
@jwt_required()
@verify_permission('view')
def get_document_versions(document_id):
    """获取文档版本历史"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if not check_document_permission(user, document):
            return jsonify({'message': '无权限访问此文档'}), 403
        
        # 仅流式文件有版本历史
        if document.file_type != 'stream':
            return jsonify({'message': '仅流式文件支持版本历史'}), 400
        
        # 获取版本历史
        versions = DocumentVersion.query.filter_by(
            document_id=document_id
        ).order_by(DocumentVersion.version_num.desc()).all()
        
        version_list = []
        for v in versions:
            version_list.append({
                'id': v.id,
                'version': v.version_num,
                'created_at': v.created_at.isoformat()
            })
        
        return jsonify({'versions': version_list})
    
    except Exception as e:
        return jsonify({'message': f'获取版本历史失败: {str(e)}'}), 500


@documents_bp.route('/<int:document_id>/versions/<int:version_id>', methods=['GET'])
@jwt_required()
@verify_permission('view')
def get_document_version(document_id, version_id):
    """获取特定版本的文档内容"""
    try:
        # 获取当前用户
        user = get_current_user()
        
        # 查找文档
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'message': '文档不存在'}), 404
        
        # 检查权限
        if not check_document_permission(user, document):
            return jsonify({'message': '无权限访问此文档'}), 403
        
        # 查找版本
        version = DocumentVersion.query.filter_by(
            id=version_id,
            document_id=document_id
        ).first()
        
        if not version:
            return jsonify({'message': '版本不存在'}), 404
        
        return jsonify({
            'version': {
                'id': version.id,
                'version': version.version_num,
                'content': version.content,
                'created_at': version.created_at.isoformat()
            }
        })
    
    except Exception as e:
        return jsonify({'message': f'获取版本内容失败: {str(e)}'}), 500


# 需要导入os模块
import os